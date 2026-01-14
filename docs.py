from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_budget_document():
    # Initialize Document
    doc = Document()

    # --- Styles Setup ---
    # Title Style
    title = doc.add_heading('WEB DEVELOPMENT PROPOSAL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle / Header Info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Prepared by: Sameer | Date: January 2, 2026')
    run.italic = True
    run.font.size = Pt(11)
    
    doc.add_paragraph() # Spacer

    # --- 1. Project Overview ---
    doc.add_heading('1. PROJECT OVERVIEW', level=1)
    p = doc.add_paragraph(
        'We are pleased to submit this proposal for the design and development of a professional website to '
        'establish your digital presence. This document outlines three distinct packages tailored to meet '
        'different levels of interactivity and customization, tailored specifically for the construction '
        'industry requirements.'
    )

    # --- 2. Development Packages ---
    doc.add_heading('2. DEVELOPMENT PACKAGES', level=1)

    # Package A
    h = doc.add_heading('Option A: The "Starter" Package', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Cost: ₹10,000 (One-time)')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0) # Dark Green
    
    doc.add_paragraph('Best for: Establishing a quick, professional online business card.', style='Intense Quote')
    
    doc.add_paragraph('Structure: Simple 5-Page Static Website (Home, About, Services, Projects, Contact).')
    
    p = doc.add_paragraph('Features:')
    p.style = 'List Bullet'
    doc.add_paragraph('Mobile responsive design', style='List Bullet 2')
    doc.add_paragraph('WhatsApp Chat Integration', style='List Bullet 2')
    doc.add_paragraph('Basic Gallery for past projects', style='List Bullet 2')
    
    p = doc.add_paragraph('Support (1 Year Included):')
    p.style = 'List Bullet'
    doc.add_paragraph('Technical fix guarantee for 12 months.', style='List Bullet 2')
    doc.add_paragraph('Minor content updates (adding/removing pictures, simple text changes).', style='List Bullet 2')

    # Package B
    h = doc.add_heading('Option B: The "Business" Package', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Cost: ₹15,000 (One-time)')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)

    doc.add_paragraph('Best for: Construction firms needing to showcase detailed portfolios.', style='Intense Quote')
    
    doc.add_paragraph('Structure: 8-10 Page Interactive Website.')
    
    p = doc.add_paragraph('Construction-Specific Features:')
    p.style = 'List Bullet'
    doc.add_paragraph('Project Filtration (Residential vs Commercial)', style='List Bullet 2')
    doc.add_paragraph('"Before & After" Sliders for construction progress', style='List Bullet 2')
    doc.add_paragraph('Downloadable Brochure (PDF) function', style='List Bullet 2')
    doc.add_paragraph('Testimonials Carousel', style='List Bullet 2')
    
    p = doc.add_paragraph('Support (1 Year Included):')
    p.style = 'List Bullet'
    doc.add_paragraph('Priority technical support.', style='List Bullet 2')
    doc.add_paragraph('Monthly addition of new project photos.', style='List Bullet 2')

    # Package C
    h = doc.add_heading('Option C: The "Premium" Package', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Cost: ₹25,000+ (Based on customization)')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)

    doc.add_paragraph('Best for: Brands wanting a luxury feel with full customization.', style='Intense Quote')
    
    doc.add_paragraph('Structure: Fully Custom Dynamic Website with Animations.')
    
    p = doc.add_paragraph('Premium Features:')
    p.style = 'List Bullet'
    doc.add_paragraph('High-End Animations (Scroll effects, hover effects)', style='List Bullet 2')
    doc.add_paragraph('Video Backgrounds (High quality loops)', style='List Bullet 2')
    doc.add_paragraph('Live Project Tracker status bar', style='List Bullet 2')
    doc.add_paragraph('Blog/News section for industry updates', style='List Bullet 2')
    
    p = doc.add_paragraph('Support (1 Year Included):')
    p.style = 'List Bullet'
    doc.add_paragraph('Dedicated 24-hour response time.', style='List Bullet 2')
    doc.add_paragraph('Full maintenance & security patches.', style='List Bullet 2')

    # --- 3. Variable Costs Table ---
    doc.add_page_break()
    doc.add_heading('3. VARIABLE COSTS (Domain & Hosting)', level=1)
    doc.add_paragraph('These costs are paid to third-party providers (e.g., GoDaddy/Hostinger) and vary based on selection.')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    
    # Table Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Est. Cost (Per Year)'
    hdr_cells[2].text = 'Notes'

    # Table Data
    data = [
        ('Domain Name', '₹800 - ₹1,200', '.com is standard, .in is cheaper'),
        ('Web Hosting', '₹2,500 - ₹4,500', 'Price varies by server speed'),
        ('Professional Email', '₹400 - ₹800 / user', 'Ex: info@yourcompany.com'),
        ('SSL Certificate', 'FREE - ₹2,000', 'Basic security is usually free')
    ]

    for item, cost, notes in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = cost
        row_cells[2].text = notes

    doc.add_paragraph() # Spacer

    # --- 4. Additional Services ---
    doc.add_heading('4. ADDITIONAL SERVICES (Optional)', level=1)
    
    addons = [
        ('Logo Design', '₹1,500', '3 concepts + vector files'),
        ('Content Writing', '₹2,000', 'Professional corporate copywriting'),
        ('Advanced SEO', '₹5,000/mo', 'Ranking for local keywords'),
        ('Social Media Setup', '₹1,000', 'FB/Insta/LinkedIn page creation')
    ]
    
    for service, price, desc in addons:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        run = p.add_run(f"{service}: ")
        run.bold = True
        p.add_run(f"{price} - {desc}")

    # --- 5. Terms & Payment ---
    doc.add_heading('5. TERMS & PAYMENT SCHEDULE', level=1)
    
    doc.add_paragraph('Support Terms:', style='Heading 3')
    doc.add_paragraph('Included support covers technical bugs and minor content updates. It does not include new page creation, major layout redesigns, or complex feature additions.', style='List Bullet')
    
    doc.add_paragraph('Payment Schedule:', style='Heading 3')
    doc.add_paragraph('50% Advance to initiate project.', style='List Number')
    doc.add_paragraph('25% Upon design approval.', style='List Number')
    doc.add_paragraph('25% Before final launch.', style='List Number')

    # --- Signature Block ---
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph('_' * 30 + '                                ' + '_' * 30)
    p.add_run('\nClient Signature                                                                    Developer Signature')

    # Save
    file_name = 'Web_Development_Budget_Proposal.docx'
    doc.save(file_name)
    print(f"Document saved as {file_name}")

if __name__ == "__main__":
    create_budget_document()
